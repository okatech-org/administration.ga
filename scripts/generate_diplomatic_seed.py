#!/usr/bin/env python3
"""
Génère `convex/migrations/seedDiplomaticTemplatesData.ts` à partir des 25
DOCX finalisés dans `TEMPLATE DOCUMENT/Modeles/`.

Extraction fidèle :
- Paragraphes du body (en respectant l'ordre body vs tables intercalées)
- Paragraphes vides (préservés comme paragraphe vide dans Tiptap)
- Alignement (left/right/center/justify, ou absent si None)
- Runs avec marks bold / italic / underline
- Tables (Tiptap table / tableRow / tableCell)

Exécution : `python3 scripts/generate_diplomatic_seed.py`
"""

import json
import os
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCX_ROOT = PROJECT_ROOT / "TEMPLATE DOCUMENT" / "Modeles"
OUTPUT = PROJECT_ROOT / "convex" / "migrations" / "seedDiplomaticTemplatesData.ts"

# ─── Liste canonique des 25 modèles ──────────────────────────────────────

TEMPLATES = [
    # (seedKey, name, category, templateType, subfolder, docxFilename)
    # `name` est aligné sur le TITRE officiel du document tel qu'il apparaît
    # dans le PDF source (accents préservés), avec un Title Case français qui
    # garde les prépositions en minuscules (« de », « à », « du »…).
    ("diplo_attestation_d_hebergement",          "Attestation d'Hébergement",               "certification", "attestation", "01 - Documents Consulaires",         "Attestation d'Hebergement.docx"),
    ("diplo_attestation_de_legalisation",        "Attestation de Légalisation",             "certification", "attestation", "01 - Documents Consulaires",         "Attestation de Legalisation.docx"),
    ("diplo_attestation_de_prise_en_charge",     "Attestation de Prise en Charge",          "certification", "attestation", "01 - Documents Consulaires",         "Attestation de Prise en Charge.docx"),
    ("diplo_certificat_de_celibat",              "Certificat de Célibat",                   "certification", "certificate", "01 - Documents Consulaires",         "Certificat de Celibat.docx"),
    ("diplo_certificat_de_nationalite",          "Certificat de Nationalité",               "certification", "certificate", "01 - Documents Consulaires",         "Certificat de Nationalite.docx"),
    ("diplo_certificat_de_residence",            "Certificat de Résidence",                 "certification", "certificate", "01 - Documents Consulaires",         "Certificat de Residence.docx"),
    ("diplo_certificat_de_vie",                  "Certificat de Vie",                       "certification", "certificate", "01 - Documents Consulaires",         "Certificat de Vie.docx"),
    ("diplo_procuration",                        "Procuration",                             "certification", "custom",      "01 - Documents Consulaires",         "Procuration.docx"),
    ("diplo_aide_memoire",                       "Aide-Mémoire",                            "notification",  "letter",      "02 - Documents Diplomatiques",       "Aide-Memoire.docx"),
    ("diplo_communique",                         "Communiqué",                              "notification",  "letter",      "02 - Documents Diplomatiques",       "Communique.docx"),
    ("diplo_demande_d_agrement",                 "Demande d'Agrément",                      "notification",  "letter",      "02 - Documents Diplomatiques",       "Demande d'Agrement.docx"),
    ("diplo_note_verbale_ministere",             "Note Verbale (Ministère)",                "notification",  "letter",      "02 - Documents Diplomatiques",       "Note Verbale (Ministere).docx"),
    ("diplo_note_verbale_missions_diplomatiques","Note Verbale (Missions Diplomatiques)",   "notification",  "letter",      "02 - Documents Diplomatiques",       "Note Verbale (Missions Diplomatiques).docx"),
    ("diplo_bordereau_d_envoi",                  "Bordereau d'Envoi",                       "notification",  "letter",      "03 - Correspondances Officielles",   "Bordereau d'Envoi.docx"),
    ("diplo_lettre_officielle_de_l_ambassadeur", "Lettre Officielle de l'Ambassadeur",      "notification",  "letter",      "03 - Correspondances Officielles",   "Lettre Officielle de l'Ambassadeur.docx"),
    ("diplo_lettre_de_felicitations",            "Lettre de Félicitations",                 "notification",  "letter",      "03 - Correspondances Officielles",   "Lettre de Felicitations.docx"),
    ("diplo_note_de_condoleances",               "Note de Condoléances",                    "notification",  "letter",      "03 - Correspondances Officielles",   "Note de Condoleances.docx"),
    ("diplo_note_de_service_interne",            "Note de Service Interne",                 "notification",  "letter",      "03 - Correspondances Officielles",   "Note de Service Interne.docx"),
    ("diplo_fiche_d_inscription_consulaire",     "Fiche d'Inscription au Registre Consulaire","registration","custom",      "04 - Identite et Voyage",            "Fiche d'Inscription Consulaire.docx"),
    ("diplo_formulaire_de_demande_de_visa",      "Formulaire de Demande de Visa",           "visa",          "custom",      "04 - Identite et Voyage",            "Formulaire de Demande de Visa.docx"),
    ("diplo_laissez_passer",                     "Laissez-Passer",                          "travel_document","custom",     "04 - Identite et Voyage",            "Laissez-Passer.docx"),
    ("diplo_attestation_sur_l_honneur",          "Attestation sur l'Honneur",               "declaration",   "attestation", "05 - Documents Notariaux et Juridiques","Attestation sur l'Honneur.docx"),
    ("diplo_certificat_de_capacite_matrimoniale","Certificat de Capacité Matrimoniale",     "certification", "certificate", "05 - Documents Notariaux et Juridiques","Certificat de Capacite Matrimoniale.docx"),
    ("diplo_certificat_de_coutume",              "Certificat de Coutume",                   "certification", "certificate", "05 - Documents Notariaux et Juridiques","Certificat de Coutume.docx"),
    ("diplo_transcription_acte_de_naissance",    "Transcription d'Acte de Naissance",       "transcript",    "certificate", "05 - Documents Notariaux et Juridiques","Transcription Acte de Naissance.docx"),
]

# ─── Helpers d'extraction ────────────────────────────────────────────────

ALIGN_MAP = {
    "LEFT": "left",
    "CENTER": "center",
    "RIGHT": "right",
    "JUSTIFY": "justify",
    "JUSTIFY_LOW": "justify",
    "JUSTIFY_HI": "justify",
    "JUSTIFY_MED": "justify",
}


def paragraph_align(p):
    if p.alignment is None:
        return None
    name = p.alignment.name
    return ALIGN_MAP.get(name)


def run_marks(run):
    """Extrait les marks Tiptap d'un run DOCX : bold / italic / underline +
    textStyle (fontFamily + fontSize). Les attributs textStyle sont fusionnés
    sur un seul mark pour respecter le schéma Tiptap (extension `TextStyle`).
    """
    marks = []
    if run.bold:
        marks.append({"type": "bold"})
    if run.italic:
        marks.append({"type": "italic"})
    if run.underline:
        marks.append({"type": "underline"})

    # Style de police (famille + taille) — exposé via le mark `textStyle`
    text_style_attrs = {}
    font = run.font
    if font.name:
        text_style_attrs["fontFamily"] = font.name
    if font.size:
        text_style_attrs["fontSize"] = int(font.size.pt)
    if text_style_attrs:
        marks.append({"type": "textStyle", "attrs": text_style_attrs})
    return marks


def paragraph_to_tiptap(p):
    """Convertit un `docx.Paragraph` en node Tiptap `paragraph`.

    - Les runs sont fusionnés en nœuds `text` avec marks.
    - Les runs consécutifs de mêmes marks sont fusionnés pour éviter le bruit.
    - Un paragraphe vide devient `{ type: "paragraph" }` (sans content).
    """
    align = paragraph_align(p)
    attrs = {"textAlign": align} if align else None

    # Collecter les runs comme (marks_fingerprint, marks, text). Le
    # fingerprint inclut les attrs (fontFamily / fontSize) pour que deux
    # runs avec la même fonte MAIS des tailles différentes ne soient pas
    # fusionnés.
    def fingerprint(marks):
        parts = []
        for m in marks:
            mtype = m["type"]
            attrs = m.get("attrs") or {}
            attrs_str = ",".join(f"{k}={attrs[k]}" for k in sorted(attrs.keys()))
            parts.append(f"{mtype}({attrs_str})")
        return tuple(parts)

    segments = []
    for run in p.runs:
        if run.text == "":
            continue
        marks = run_marks(run)
        segments.append((fingerprint(marks), marks, run.text))

    # Fusion des runs consécutifs avec les mêmes marks
    merged = []
    for key, marks, text in segments:
        if merged and merged[-1][0] == key:
            merged[-1] = (key, marks, merged[-1][2] + text)
        else:
            merged.append((key, marks, text))

    content = []
    for _, marks, text in merged:
        node = {"type": "text", "text": text}
        if marks:
            node["marks"] = marks
        content.append(node)

    out = {"type": "paragraph"}
    if attrs:
        out["attrs"] = attrs
    if content:
        out["content"] = content
    return out


def table_to_tiptap(tbl):
    """Convertit une `docx.Table` en node Tiptap `table` (row / tableCell).

    Chaque cellule est un ensemble de paragraphes Tiptap.
    """
    rows = []
    for row in tbl.rows:
        cells = []
        for cell in row.cells:
            cell_content = [paragraph_to_tiptap(p) for p in cell.paragraphs]
            # Au moins un paragraphe vide si la cellule n'en a aucun
            if not cell_content:
                cell_content = [{"type": "paragraph"}]
            cells.append(
                {
                    "type": "tableCell",
                    "attrs": {"colspan": 1, "rowspan": 1},
                    "content": cell_content,
                }
            )
        rows.append({"type": "tableRow", "content": cells})
    return {"type": "table", "content": rows}


def iter_body_blocks(doc):
    """Itère sur les blocs top-level dans l'ordre (paragraphes + tables).

    `python-docx` expose `doc.paragraphs` et `doc.tables` séparément, ce
    qui perd l'ordre. On passe donc par le XML body pour respecter la
    séquence réelle.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            # Retrouve l'objet Paragraph correspondant
            for p in doc.paragraphs:
                if p._element is child:
                    yield ("p", p)
                    break
        elif child.tag == qn("w:tbl"):
            for t in doc.tables:
                if t._element is child:
                    yield ("t", t)
                    break


def docx_to_tiptap(path):
    doc = docx.Document(path)
    content = []
    for kind, block in iter_body_blocks(doc):
        if kind == "p":
            content.append(paragraph_to_tiptap(block))
        elif kind == "t":
            content.append(table_to_tiptap(block))
    # Supprime les paragraphes vides strictement en tête et en queue pour
    # éviter les marges superflues au rendu — on préserve ceux du milieu.
    while content and is_empty_paragraph(content[0]):
        content.pop(0)
    while content and is_empty_paragraph(content[-1]):
        content.pop()
    return {"type": "doc", "content": content}


def is_empty_paragraph(node):
    return (
        node.get("type") == "paragraph"
        and not node.get("content")
        and not node.get("attrs")
    )


# ─── Génération du fichier TS ────────────────────────────────────────────


def escape_ts_backtick(s):
    """Échappe une chaîne pour une template literal TS (backticks)."""
    return s.replace("\\", "\\\\").replace("`", "\\`").replace("${", "\\${")


def generate_ts_file(templates):
    blocks = []
    for seed_key, name, category, template_type, subfolder, filename in templates:
        docx_path = DOCX_ROOT / subfolder / filename
        if not docx_path.exists():
            print(f"  MISSING: {docx_path}", file=sys.stderr)
            continue
        content = docx_to_tiptap(docx_path)
        content_json = json.dumps(content, ensure_ascii=False, separators=(", ", ": "))

        blocks.append(
            "\t{\n"
            f'\t\tseedKey: "{seed_key}",\n'
            f"\t\tname: `{escape_ts_backtick(name)}`,\n"
            f'\t\tcategory: "{category}",\n'
            f'\t\ttemplateType: "{template_type}",\n'
            f"\t\tsubfolder: `{escape_ts_backtick(subfolder)}`,\n"
            f"\t\tcontent: {content_json},\n"
            "\t},"
        )

    header = '''/**
 * Données des 25 modèles diplomatiques/consulaires du Gabon à Madrid.
 * Généré depuis les 25 DOCX finalisés dans TEMPLATE DOCUMENT/Modeles/.
 * Utilisé par `seedDefaultTemplates` (idempotent — key = name.fr).
 *
 * ⚠️ Ne pas éditer à la main : régénérer via
 *    python3 scripts/generate_diplomatic_seed.py
 *
 * Conserve fidèlement :
 *   - paragraphes vides (sauts de ligne visuels)
 *   - alignements (left / right / center / justify)
 *   - marks runs (bold / italic / underline)
 *   - tables (titre centré 1x1, grilles de saisie 2 colonnes)
 */

export type DiplomaticTemplateSeed = {
\tseedKey: string;
\tname: string;
\tcategory:
\t\t| "identity"
\t\t| "passport"
\t\t| "civil_status"
\t\t| "visa"
\t\t| "certification"
\t\t| "transcript"
\t\t| "registration"
\t\t| "notification"
\t\t| "assistance"
\t\t| "travel_document"
\t\t| "declaration"
\t\t| "other";
\ttemplateType: "certificate" | "attestation" | "receipt" | "letter" | "custom";
\tsubfolder: string;
\tcontent: unknown;
};

export const DIPLOMATIC_TEMPLATES: DiplomaticTemplateSeed[] = [
'''
    footer = "\n];\n"
    return header + "\n".join(blocks) + footer


def main():
    print(f"DOCX root: {DOCX_ROOT}")
    print(f"Output: {OUTPUT}")
    ts = generate_ts_file(TEMPLATES)
    OUTPUT.write_text(ts, encoding="utf-8")
    total_bytes = len(ts.encode("utf-8"))
    print(f"✔ Wrote {len(TEMPLATES)} templates ({total_bytes} bytes)")


if __name__ == "__main__":
    main()
